"""Binary document -> Markdown conversion with an on-disk cache.

PDF is converted via pymupdf4llm; DOCX is converted via python-docx. The cache
path mirrors the content hash of the source so repeated indexing of an
unchanged file skips the expensive conversion step.

Cache layout:
    $MFS_HOME/converted/<hash[:2]>/<hash>.md

LRU eviction uses ``cache.max_size_mb`` and deletes the oldest converted files
when the cache grows past the configured cap.
"""

from __future__ import annotations

import hashlib
import os
from pathlib import Path

from ..config import ensure_mfs_home, load_config


def is_convertible(extension: str) -> bool:
    return extension.lower() in {".pdf", ".docx"}


def _cache_path(file_hash: str) -> Path:
    home = ensure_mfs_home()
    return home / "converted" / file_hash[:2] / f"{file_hash}.md"


def _hash_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1 << 16), b""):
            h.update(chunk)
    return h.hexdigest()


def convert_to_markdown(path: Path) -> str:
    """Return Markdown text for *path*, using the on-disk cache when possible.

    Raises ``RuntimeError`` if the format isn't supported or the converter
    library is missing.
    """
    ext = path.suffix.lower()
    if ext not in {".pdf", ".docx"}:
        raise RuntimeError(f"no converter registered for {ext}")

    file_hash = _hash_file(path)
    cache = _cache_path(file_hash)
    if cache.exists():
        try:
            text = cache.read_text(encoding="utf-8")
            _touch(cache)
            return text
        except OSError:
            pass  # fall through and re-convert

    if ext == ".pdf":
        markdown = _convert_pdf(path)
    else:
        markdown = _convert_docx(path)

    try:
        cache.parent.mkdir(parents=True, exist_ok=True)
        cache.write_text(markdown, encoding="utf-8")
        _enforce_cache_limit()
    except OSError:
        # Cache is best-effort; ignore write failures (disk full, permissions).
        pass
    return markdown


def _convert_pdf(path: Path) -> str:
    try:
        import pymupdf4llm  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "PDF conversion requires pymupdf4llm. "
            "Install with: uv add pymupdf4llm"
        ) from exc

    try:
        return pymupdf4llm.to_markdown(str(path))
    except Exception as exc:
        raise RuntimeError(f"pymupdf4llm failed on {path}: {exc}") from exc


def _convert_docx(path: Path) -> str:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "DOCX conversion requires python-docx. Install with: uv add python-docx"
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"python-docx failed on {path}: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
        lines.append("")

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if rows:
            width = max(len(r) for r in rows)
            normalized = [r + [""] * (width - len(r)) for r in rows]
            lines.append("| " + " | ".join(normalized[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            for row in normalized[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")

    return "\n".join(lines).strip() + "\n"


def _heading_level(style_name: str) -> int:
    parts = style_name.split()
    for part in reversed(parts):
        if part.isdigit():
            return min(6, max(1, int(part)))
    return 1


def _touch(path: Path) -> None:
    try:
        os.utime(path, None)
    except OSError:
        pass


def _enforce_cache_limit() -> None:
    try:
        max_size_mb = int(load_config().cache.max_size_mb)
    except Exception:
        max_size_mb = 500
    if max_size_mb <= 0:
        return

    root = ensure_mfs_home() / "converted"
    limit = max_size_mb * 1024 * 1024
    files: list[tuple[float, int, Path]] = []
    total = 0
    for item in root.rglob("*.md"):
        try:
            st = item.stat()
        except OSError:
            continue
        total += st.st_size
        files.append((st.st_atime, st.st_size, item))
    if total <= limit:
        return

    for _atime, size, item in sorted(files, key=lambda x: x[0]):
        try:
            item.unlink()
        except OSError:
            continue
        total -= size
        _remove_empty_parents(item.parent, root)
        if total <= limit:
            break


def _remove_empty_parents(path: Path, stop: Path) -> None:
    while path != stop and path.is_relative_to(stop):
        try:
            path.rmdir()
        except OSError:
            return
        path = path.parent
