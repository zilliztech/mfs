"""Regression tests for the e2e fixes (Bugs 1-6 from the 2026-04-20 audit).

Bug 1 is a pyproject.toml-only change (verified by `uv sync --extra onnx`);
this module covers Bugs 2, 3, 5, and 6 in code. Bug 4 (Milvus-lock retry
+ watch-cycle close) is behavioral and exercised indirectly via the
``_looks_like_lock_error`` helper.
"""

from __future__ import annotations

from click.testing import CliRunner

from mfs.cli import (
    _find_image_files,
    _looks_like_lock_error,
)
from mfs.cli_config import config_group
from mfs.ingest import converter as converter_mod
from mfs.ingest.converter import convert_to_markdown, is_convertible


# ---------------------------------------------------------------- Bug 2: PDF


def test_is_convertible_pdf_and_docx():
    """PDF and DOCX are converted to Markdown; EPUB stays out of the MVP."""
    assert is_convertible(".pdf")
    assert is_convertible(".PDF")
    assert is_convertible(".docx")
    assert is_convertible(".DOCX")
    assert not is_convertible(".epub")
    assert not is_convertible(".md")


def test_convert_to_markdown_caches_output(tmp_path, mfs_home, monkeypatch):
    """First conversion writes cache; a second call with unchanged bytes
    reads the cache file (no converter re-invocation)."""
    pdf = tmp_path / "tiny.pdf"
    import pymupdf  # type: ignore[import-not-found]

    doc = pymupdf.Document()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello MFS PDF chunking")
    doc.save(str(pdf))
    doc.close()

    md1 = convert_to_markdown(pdf)
    assert "Hello MFS PDF chunking" in md1

    # Exactly one cache artifact written under ~/.mfs/converted/<hh>/<hash>.md.
    cache_dir = mfs_home / "converted"
    cache_files = list(cache_dir.rglob("*.md"))
    assert len(cache_files) == 1
    assert "Hello MFS PDF chunking" in cache_files[0].read_text()

    # Second call with the same bytes must hit the cache. We prove this by
    # monkey-patching the converter so any re-entry raises loudly.
    import pymupdf4llm  # type: ignore[import-not-found]

    def _boom(*_a, **_kw):  # pragma: no cover
        raise AssertionError("pymupdf4llm should not be re-invoked on cache hit")

    monkeypatch.setattr(pymupdf4llm, "to_markdown", _boom)
    md2 = convert_to_markdown(pdf)
    assert md2 == md1


def test_convert_to_markdown_supports_docx(tmp_path, mfs_home):
    docx_path = tmp_path / "guide.docx"
    import docx

    document = docx.Document()
    document.add_heading("DOCX Guide", level=1)
    document.add_paragraph("Token expiration is handled here.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Error"
    table.cell(1, 1).text = "ERR_TOKEN_EXPIRED"
    document.save(str(docx_path))

    markdown = convert_to_markdown(docx_path)
    assert "# DOCX Guide" in markdown
    assert "Token expiration is handled here." in markdown
    assert "| Key | Value |" in markdown
    assert "ERR_TOKEN_EXPIRED" in markdown


def test_converted_cache_lru_eviction(mfs_home, monkeypatch):
    class _Cache:
        max_size_mb = 1

    class _Config:
        cache = _Cache()

    monkeypatch.setattr(converter_mod, "load_config", lambda: _Config())
    cache_root = mfs_home / "converted"
    old_dir = cache_root / "aa"
    new_dir = cache_root / "bb"
    old_dir.mkdir(parents=True)
    new_dir.mkdir(parents=True)
    old_file = old_dir / "old.md"
    new_file = new_dir / "new.md"
    old_file.write_bytes(b"a" * 700_000)
    new_file.write_bytes(b"b" * 700_000)
    old_time = 1_700_000_000
    new_time = 1_700_000_100
    import os

    os.utime(old_file, (old_time, old_time))
    os.utime(new_file, (new_time, new_time))

    converter_mod._enforce_cache_limit()

    assert not old_file.exists()
    assert new_file.exists()


def test_convertible_constants_include_docx_exclude_epub():
    """INDEXED_EXTENSIONS must not silently index unsupported binary docs."""
    from mfs import constants as C
    assert ".pdf" in C.INDEXED_EXTENSIONS
    assert ".docx" in C.INDEXED_EXTENSIONS
    assert ".epub" not in C.INDEXED_EXTENSIONS


# ---------------------------------------------------------------- Bug 3: watch-delete


def test_run_add_once_processes_deletion_on_empty_dir(mfs_home, tmp_path, monkeypatch):
    """Deleting the last file under a watched root must propagate to Milvus."""
    from mfs.cli import _run_add_once
    from mfs.config import load_config
    from mfs.embedder.openai import OpenAIEmbedding
    from mfs.ingest.scanner import Scanner
    from mfs.store import ChunkRecord, MilvusStore

    work = tmp_path / "wt"
    work.mkdir()
    target = work / "a.md"
    target.write_text("# A\n\nhello\n", encoding="utf-8")

    # Stub the embedder so we don't need the OpenAI API during test.
    class _StubEmbedder:
        model_name = "stub"
        dimension = 8
        batch_size = 4

        def embed(self, texts):
            return [[0.0] * 8 for _ in texts]

    monkeypatch.setattr("mfs.cli._build_embedder", lambda _cfg: _StubEmbedder())
    monkeypatch.setattr(
        OpenAIEmbedding, "__init__", lambda self, *a, **kw: None, raising=False
    )

    config = load_config()
    store = MilvusStore(config.milvus, 8)
    store.connect()
    # Seed one indexed body chunk for a.md so there's state to delete.
    store.insert_chunks([
        ChunkRecord(
            id="seed0001",
            source=str(target.resolve()),
            parent_dir=str(work.resolve()),
            chunk_index=0,
            start_line=1,
            end_line=2,
            chunk_text="hello",
            dense_vector=[0.0] * 8,
            content_type="markdown",
            file_hash="deadbeef",
            is_dir=False,
            embed_status="complete",
            metadata={},
            account_id=config.milvus.account_id,
        )
    ])
    assert store.get_indexed_files(str(work.resolve()))

    scanner = Scanner(config)
    # Now delete the file and re-run add. With the fix, the empty-file-set path
    # still reaches compute_diff, notices the deletion, and drops the row.
    target.unlink()

    _run_add_once(
        [work.resolve()],
        force=False,
        sync_mode=True,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )
    assert store.get_indexed_files(str(work.resolve())) == {}


# ---------------------------------------------------------------- Bug 4: lock error


def test_looks_like_lock_error_recognises_milvus_lite_strings():
    """The heuristic catches the canonical Milvus Lite lock messages."""
    assert _looks_like_lock_error(
        RuntimeError("file is opened by another program")
    )
    assert _looks_like_lock_error(
        RuntimeError("database is locked")
    )
    assert _looks_like_lock_error(
        RuntimeError("Resource temporarily unavailable")
    )
    assert not _looks_like_lock_error(RuntimeError("some other failure"))


# ---------------------------------------------------------------- Bug 5: --describe single image


def test_find_image_files_accepts_single_file(tmp_path):
    """`_find_image_files` must return a single-file path ending in an image ext."""
    img = tmp_path / "pic.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
    out = _find_image_files([img])
    assert out == [img]

    other = tmp_path / "notes.md"
    other.write_text("# hi", encoding="utf-8")
    assert _find_image_files([other]) == []


# ---------------------------------------------------------------- Bug 6: config set auto-model


def test_config_set_embedding_provider_realigns_model(mfs_home):
    runner = CliRunner()
    # Seed an OpenAI-default config.toml.
    result = runner.invoke(config_group, ["init", "--force"])
    assert result.exit_code == 0, result.output
    runner.invoke(config_group, ["set", "embedding.provider", "openai"])
    runner.invoke(config_group, ["set", "embedding.model", "text-embedding-3-small"])

    # Switching to onnx must auto-set the model to the onnx default.
    result = runner.invoke(config_group, ["set", "embedding.provider", "onnx"])
    assert result.exit_code == 0, result.output
    assert "gpahal/bge-m3-onnx-int8" in result.output

    get = runner.invoke(config_group, ["get", "embedding.model"])
    assert get.output.strip() == "gpahal/bge-m3-onnx-int8"


def test_config_set_preserves_custom_model(mfs_home):
    """If the user set a bespoke model string, switching provider doesn't clobber it."""
    runner = CliRunner()
    runner.invoke(config_group, ["init", "--force"])
    # Custom model string (not any of the known provider defaults).
    runner.invoke(config_group, ["set", "embedding.model", "my-custom-embedder:v1"])

    result = runner.invoke(config_group, ["set", "embedding.provider", "onnx"])
    assert result.exit_code == 0
    # The note shouldn't be printed because the model is custom.
    assert "gpahal/bge-m3-onnx-int8" not in result.output

    get = runner.invoke(config_group, ["get", "embedding.model"])
    assert get.output.strip() == "my-custom-embedder:v1"


def test_config_set_llm_provider_realigns_model(mfs_home):
    runner = CliRunner()
    runner.invoke(config_group, ["init", "--force"])
    runner.invoke(config_group, ["set", "llm.provider", "openai"])
    runner.invoke(config_group, ["set", "llm.model", "gpt-4o-mini"])

    result = runner.invoke(config_group, ["set", "llm.provider", "anthropic"])
    assert result.exit_code == 0
    assert "claude-3-5-haiku-latest" in result.output

    get = runner.invoke(config_group, ["get", "llm.model"])
    assert get.output.strip() == "claude-3-5-haiku-latest"


# ---------------------------------------------------------------- grep -C separator


def _line_numbers_in_gutter(lines: list[str]) -> list[int]:
    """Parse left-gutter line numbers from formatted grep/search output.

    Gutter rows look like ``"   7  content"`` (right-aligned number, then
    exactly two spaces before the payload). Headers and ``--`` separators
    don't match this shape.
    """
    import re as _re
    nums: list[int] = []
    for ln in lines:
        m = _re.match(r"^\s*(\d+) {2}", ln)
        if m:
            nums.append(int(m.group(1)))
    return nums


def test_format_grep_results_emits_separator_between_gaps():
    """Non-contiguous line blocks within the same file get a '--' separator."""
    from mfs.output.display import format_grep_results
    from mfs.search.searcher import GrepMatch

    src = "/tmp/demo.md"
    # Match at line 1 with -C 2 -> context reaches line 3.
    m1 = GrepMatch(source=src, line_number=1, line_text="hit one",
                   context_before=[], context_after=["line2", "line3"])
    # Match at line 5 with -C 2 -> context covers lines 3..7.
    # Since line 3 is already "touching" line 3 above (same line), blocks should
    # merge with no separator.
    m2 = GrepMatch(source=src, line_number=5, line_text="hit two",
                   context_before=["line3", "line4"], context_after=["line6", "line7"])
    # Match at line 20 with -C 2 -> context 18..22. Gap from line 7 to 18.
    m3 = GrepMatch(source=src, line_number=20, line_text="hit three",
                   context_before=["line18", "line19"], context_after=["line21", "line22"])

    out = format_grep_results([m1, m2, m3])
    lines = out.splitlines()

    # Exactly one '--' separator: between the merged (1..7) block and the (18..22) block.
    assert lines.count("--") == 1
    sep_idx = lines.index("--")
    before_nums = _line_numbers_in_gutter(lines[:sep_idx])
    after_nums = _line_numbers_in_gutter(lines[sep_idx + 1:])
    assert 7 in before_nums and 18 in after_nums


def test_format_grep_results_no_separator_when_blocks_touch():
    """Adjacent context windows (end of one equals start-1 of next) must not get '--'."""
    from mfs.output.display import format_grep_results
    from mfs.search.searcher import GrepMatch

    src = "/tmp/demo.md"
    # Line 5 with -C 2 -> 3..7. Line 8 with -C 2 -> 6..10. These windows overlap.
    m1 = GrepMatch(source=src, line_number=5, line_text="five",
                   context_before=["line3", "line4"], context_after=["line6", "line7"])
    m2 = GrepMatch(source=src, line_number=8, line_text="eight",
                   context_before=["line6", "line7"], context_after=["line9", "line10"])
    out = format_grep_results([m1, m2])
    assert "--" not in out.splitlines()


# ---------------------------------------------------------------- Bug: sync mode O(N²) / Bug: config set list values


def test_add_sync_does_not_use_queue_json(mfs_home, tmp_path, monkeypatch):
    """`mfs add --sync` must stream tasks directly to Milvus without round-
    tripping through queue.json. At 100 files we'd previously pay O(N²) I/O
    on the atomic queue-file rewrite (minutes per chunk at 10K chunks); the
    streaming path should finish quickly and leave the queue untouched.
    """
    import time

    from mfs.cli import _run_add_once
    from mfs.config import load_config
    from mfs.embedder.openai import OpenAIEmbedding
    from mfs.ingest.scanner import Scanner
    from mfs.store import MilvusStore

    work = tmp_path / "corpus"
    work.mkdir()
    # 100 fake markdown files, ~1KB each, with a heading so they chunk at
    # least once through the markdown splitter.
    body = "# File {i}\n\n" + ("Lorem ipsum dolor sit amet. " * 20) + "\n"
    for i in range(100):
        (work / f"note_{i:03d}.md").write_text(
            body.format(i=i), encoding="utf-8"
        )

    class _StubEmbedder:
        model_name = "stub"
        dimension = 8
        batch_size = 16

        def embed(self, texts):
            return [[0.0] * 8 for _ in texts]

    monkeypatch.setattr("mfs.cli._build_embedder", lambda _cfg: _StubEmbedder())
    monkeypatch.setattr(
        OpenAIEmbedding, "__init__", lambda self, *a, **kw: None, raising=False
    )

    config = load_config()
    store = MilvusStore(config.milvus, 8)
    store.connect()
    scanner = Scanner(config)

    start = time.monotonic()
    _run_add_once(
        [work.resolve()],
        force=False,
        sync_mode=True,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )
    elapsed = time.monotonic() - start

    # With a stub embedder, 100 tiny markdown files must finish well under
    # 60 s. (Pre-fix the queue rewrite alone blew past this on 1K+ files; we
    # keep the bound loose so a cold CI still comfortably passes.)
    assert elapsed < 60.0, f"sync-mode took {elapsed:.1f}s — queue round-trip likely regressed"

    # All 100 files should be represented in Milvus.
    indexed = store.get_indexed_files(str(work.resolve()))
    assert len(indexed) == 100, f"expected 100 indexed files, got {len(indexed)}"

    # Sync path must not create queue.json (or leave it non-empty).
    queue_path = mfs_home / "queue.json"
    if queue_path.exists():
        import json
        data = json.loads(queue_path.read_text(encoding="utf-8"))
        assert data.get("tasks") in (None, []), (
            f"sync-mode should not populate queue.json; found {len(data['tasks'])} tasks"
        )


def test_add_sync_small_corpus_still_works(mfs_home, tmp_path, monkeypatch):
    """The simple case (a handful of files) must keep working under the
    streaming sync path — no progress-bar threshold edge-case regressions.
    """
    from mfs.cli import _run_add_once
    from mfs.config import load_config
    from mfs.embedder.openai import OpenAIEmbedding
    from mfs.ingest.scanner import Scanner
    from mfs.store import MilvusStore

    work = tmp_path / "small"
    work.mkdir()
    for i in range(3):
        (work / f"f{i}.md").write_text(f"# F{i}\n\nhello {i}\n", encoding="utf-8")

    class _StubEmbedder:
        model_name = "stub"
        dimension = 8
        batch_size = 16

        def embed(self, texts):
            return [[0.0] * 8 for _ in texts]

    monkeypatch.setattr("mfs.cli._build_embedder", lambda _cfg: _StubEmbedder())
    monkeypatch.setattr(
        OpenAIEmbedding, "__init__", lambda self, *a, **kw: None, raising=False
    )

    config = load_config()
    store = MilvusStore(config.milvus, 8)
    store.connect()
    scanner = Scanner(config)

    _run_add_once(
        [work.resolve()],
        force=False,
        sync_mode=True,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )
    indexed = store.get_indexed_files(str(work.resolve()))
    assert len(indexed) == 3


def test_async_queue_uses_chunk_refs_and_priority_order(mfs_home, tmp_path, monkeypatch):
    """Async add should keep queue.json light and sort high-value files first."""
    import json

    from mfs.cli import _run_add_once
    from mfs.config import load_config
    from mfs.ingest.scanner import Scanner
    from mfs.store import MilvusStore

    work = tmp_path / "priority"
    (work / "src").mkdir(parents=True)
    (work / "tests").mkdir()
    (work / "README.md").write_text("# Project\n\noverview\n", encoding="utf-8")
    (work / "src" / "main.py").write_text("def main():\n    return 1\n", encoding="utf-8")
    (work / "tests" / "test_main.py").write_text("def test_main():\n    assert True\n", encoding="utf-8")

    class _StubEmbedder:
        model_name = "stub"
        dimension = 8
        batch_size = 16

    monkeypatch.setattr("mfs.cli.Worker.ensure_running", lambda self: None)

    config = load_config()
    store = MilvusStore(config.milvus, 8)
    store.connect()
    scanner = Scanner(config)

    _run_add_once(
        [work.resolve()],
        force=False,
        sync_mode=False,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )

    data = json.loads((mfs_home / "queue.json").read_text(encoding="utf-8"))
    tasks = data["tasks"]
    assert tasks
    assert all(t["task_type"] == "embed_ref" for t in tasks)
    assert all(t["chunk_text"] == "" for t in tasks)
    assert all(t["content_hash"] for t in tasks)
    assert tasks[0]["source"].endswith("README.md")
    assert any("/src/" in t["source"] for t in tasks[:2])


def test_worker_restores_embed_ref_before_insert(tmp_path):
    """The worker can recover chunk text from a lightweight queue reference."""
    import logging

    from mfs.ingest.chunker import chunk_file, generate_chunk_id, hash_text
    from mfs.ingest.queue import QueueTask
    from mfs.ingest.worker import process_batch

    source = tmp_path / "note.md"
    text = "# Note\n\nimportant body\n"
    source.write_text(text, encoding="utf-8")
    file_hash = __import__("hashlib").sha256(source.read_bytes()).hexdigest()
    chunk = chunk_file(source, text, ".md")[0]
    content_hash = hash_text(chunk.text)
    chunk_id = generate_chunk_id(
        str(source), chunk.start_line, chunk.end_line, content_hash, "stub"
    )

    task = QueueTask(
        chunk_id=chunk_id,
        source=str(source),
        parent_dir=str(tmp_path),
        chunk_text="",
        chunk_index=chunk.chunk_index,
        start_line=chunk.start_line,
        end_line=chunk.end_line,
        content_type=chunk.content_type,
        file_hash=file_hash,
        is_dir=False,
        metadata={},
        account_id="default",
        content_hash=content_hash,
        task_type="embed_ref",
    )

    class _StubEmbedder:
        def embed(self, texts):
            assert texts == [chunk.text]
            return [[0.0] * 4]

    class _StubStore:
        records = None

        def insert_chunks(self, records):
            self.records = records

    store = _StubStore()
    count = process_batch([task], _StubEmbedder(), store, logging.getLogger("test"))
    assert count == 1
    assert store.records[0].chunk_text == chunk.text


def test_modified_file_queues_only_new_chunk_refs(mfs_home, tmp_path, monkeypatch):
    """Modified files keep unchanged chunks and queue only new chunk refs."""
    import json

    from mfs.cli import _run_add_once
    from mfs.config import load_config
    from mfs.ingest.scanner import Scanner
    from mfs.store import MilvusStore

    work = tmp_path / "diff"
    work.mkdir()
    doc = work / "doc.md"
    doc.write_text("# A\n\naaa\n\n# B\n\nbbb\n", encoding="utf-8")

    class _StubEmbedder:
        model_name = "stub"
        dimension = 8
        batch_size = 16

        def embed(self, texts):
            return [[0.0] * 8 for _ in texts]

    config = load_config()
    store = MilvusStore(config.milvus, 8)
    store.connect()
    scanner = Scanner(config)

    _run_add_once(
        [work.resolve()],
        force=False,
        sync_mode=True,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )
    assert len(store.get_body_chunk_ids(str(doc.resolve()))) == 2

    doc.write_text("# A\n\naaa\n\n# B\n\nchanged\n", encoding="utf-8")
    monkeypatch.setattr("mfs.cli.Worker.ensure_running", lambda self: None)
    _run_add_once(
        [work.resolve()],
        force=True,
        sync_mode=False,
        quiet=True,
        config=config,
        embedder=_StubEmbedder(),
        store=store,
        scanner=scanner,
    )

    data = json.loads((mfs_home / "queue.json").read_text(encoding="utf-8"))
    assert len(data["tasks"]) == 1
    assert data["tasks"][0]["task_type"] == "embed_ref"
    assert data["tasks"][0]["chunk_text"] == ""
    assert len(store.get_body_chunk_ids(str(doc.resolve()))) == 1
    assert store.get_indexed_files(str(work.resolve()))[str(doc.resolve())] == scanner.compute_file_hash(doc)


def test_config_set_accepts_json_list_values(mfs_home):
    """`mfs config set indexing.include_extensions '["py","md"]'` must parse
    the JSON array, not shove the raw string into TOML as a csv."""
    try:
        import tomllib as _tomllib  # type: ignore[attr-defined]
    except ModuleNotFoundError:  # Python 3.10
        import tomli as _tomllib  # type: ignore[no-redef]

    from mfs.config import config_path

    runner = CliRunner()
    result = runner.invoke(config_group, ["init", "--force"])
    assert result.exit_code == 0, result.output

    result = runner.invoke(
        config_group,
        ["set", "indexing.include_extensions", '["py","md","rst"]'],
    )
    assert result.exit_code == 0, result.output

    get = runner.invoke(config_group, ["get", "indexing.include_extensions"])
    assert get.output.strip() == "py,md,rst"

    # The raw TOML file should contain a proper array of three strings.
    with open(config_path(), "rb") as fh:
        data = _tomllib.load(fh)
    assert data["indexing"]["include_extensions"] == ["py", "md", "rst"]


def test_config_set_still_accepts_csv_for_lists(mfs_home):
    """Existing CSV form must keep working so older scripts don't break."""
    runner = CliRunner()
    runner.invoke(config_group, ["init", "--force"])
    result = runner.invoke(
        config_group,
        ["set", "indexing.exclude_extensions", "log,bin"],
    )
    assert result.exit_code == 0, result.output
    get = runner.invoke(config_group, ["get", "indexing.exclude_extensions"])
    assert get.output.strip() == "log,bin"


def test_config_set_rejects_malformed_json_array(mfs_home):
    """Leading `[` signals JSON intent — fail with a clear error if it doesn't parse."""
    runner = CliRunner()
    runner.invoke(config_group, ["init", "--force"])
    result = runner.invoke(
        config_group,
        ["set", "indexing.include_extensions", "[not-json,"],
    )
    assert result.exit_code == 2
    assert "invalid" in result.output.lower() or "invalid" in (result.stderr or "").lower()


def test_format_grep_results_separates_sources_with_blank_and_header():
    """When matches cross a file boundary, a blank line + new path header appears."""
    from mfs.output.display import format_grep_results
    from mfs.search.searcher import GrepMatch

    m1 = GrepMatch(source="/tmp/a.md", line_number=1, line_text="a-hit",
                   context_before=[], context_after=[])
    m2 = GrepMatch(source="/tmp/b.md", line_number=1, line_text="b-hit",
                   context_before=[], context_after=[])
    out = format_grep_results([m1, m2])
    lines = out.splitlines()
    # Both file paths appear as standalone headers.
    assert any(ln.endswith("/tmp/a.md") or ln == "/tmp/a.md" or "a.md" in ln
               for ln in lines)
    assert any(ln.endswith("/tmp/b.md") or ln == "/tmp/b.md" or "b.md" in ln
               for ln in lines)
    # A blank line separates the two file groups.
    assert "" in lines
